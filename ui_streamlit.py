import os
import sys
import io
import re
import base64
from openai import OpenAI
from dotenv import load_dotenv
import subprocess
from pathlib import Path

import pandas as pd
import yaml
import streamlit as st

st.session_state.setdefault("brief_docx", None)
st.session_state.setdefault("brief_txt", None)
st.session_state.setdefault("brief_brand", None)
st.session_state.setdefault("authed", False)
st.session_state.setdefault("app_pwd", "")

load_dotenv()

# allow Streamlit secrets (e.g., cloud deploy) to populate env for OpenAI SDK
if "OPENAI_API_KEY" not in os.environ and "OPENAI_API_KEY" in st.secrets:
	os.environ["OPENAI_API_KEY"] = st.secrets["OPENAI_API_KEY"]
# pass Postgres URL from Streamlit secrets to env for subprocesses
if "POSTGRES_URL" not in os.environ and "POSTGRES_URL" in st.secrets:
	os.environ["POSTGRES_URL"] = st.secrets["POSTGRES_URL"]

APP_PASSWORD = os.getenv("APP_PASSWORD") or st.secrets.get("APP_PASSWORD", "")

PROJECT_ROOT = Path(__file__).resolve().parent
ORCH = PROJECT_ROOT / "src" / "orchestrator.py"


def generate_ooda_briefing(brand: str, items_text: str) -> str:
	client = OpenAI()

	prompt = f"""
You are an AI reputation analyst. Output only the ACT section of an OODA loop briefing
for the brand shown below. Do NOT include OBSERVE, ORIENT, or DECIDE — they are
already displayed elsewhere.

Goal: cluster all items (max 30) into 2–5 thematic buckets and provide
business-ready guidance for each cluster.

Format (Markdown):
- Do NOT add any standalone heading (the UI already shows ACT).
- For each cluster, start with a level-4 heading containing only the macro theme,
  then a bullet list with:
    * Articles: <count> items
    * What it is: <1–2 sentence summary with one Markdown link [title](url) if possible>
    * Risks: <1–2 bullets of business/reputational risk>
    * Recommended action: <clear mitigation / containment / comms plan>
- Keep total length ~200 words.
- Use Markdown links: [title](url)
- If an issue is about counterfeit already seized by authorities, do NOT suggest legal escalation.

BRAND: {brand}

DATA (each item):
{items_text}
"""

	resp = client.chat.completions.create(
		model="gpt-5-mini",
		messages=[
			{"role": "system", "content": "You are a strategic reputation monitoring assistant."},
			{"role": "user", "content": prompt},
		],
	)

	return resp.choices[0].message.content


def briefing_to_docx(brand: str, briefing_text: str) -> bytes | None:
	try:
		from docx import Document  # type: ignore
		from docx.oxml import OxmlElement  # type: ignore
		from docx.oxml.ns import qn  # type: ignore
		from docx.shared import Pt, RGBColor  # type: ignore
	except ImportError:
		return None

	def add_hyperlink(paragraph, text, url):
		# Adapted helper to create hyperlink inline
		part = paragraph.part
		r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
		hyperlink = OxmlElement('w:hyperlink')
		hyperlink.set(qn('r:id'), r_id)

		new_run = OxmlElement('w:r')
		r_pr = OxmlElement('w:rPr')
		r_style = OxmlElement('w:rStyle')
		r_style.set(qn('w:val'), 'Hyperlink')
		r_pr.append(r_style)
		new_run.append(r_pr)
		text_elm = OxmlElement('w:t')
		text_elm.text = text
		new_run.append(text_elm)
		hyperlink.append(new_run)
		paragraph._p.append(hyperlink)

	def add_text_with_links(paragraph, line: str):
		# parse links, bold **text**, italics *text*
		pattern = re.compile(r'\[([^\]]+)\]\(([^)]+)\)|(\*\*([^*]+)\*\*)|(\*([^*]+)\*)')
		pos = 0
		for m in pattern.finditer(line):
			if m.start() > pos:
				paragraph.add_run(line[pos:m.start()])
			if m.group(1) and m.group(2):
				add_hyperlink(paragraph, m.group(1), m.group(2))
			elif m.group(4):
				run = paragraph.add_run(m.group(4))
				run.bold = True
			elif m.group(6):
				run = paragraph.add_run(m.group(6))
				run.italic = True
			pos = m.end()
		if pos < len(line):
			paragraph.add_run(line[pos:])

	bio = io.BytesIO()
	doc = Document()
	# base style
	style = doc.styles["Normal"]
	style.font.name = "Calibri"
	style.font.size = Pt(11)

	def heading(text, level=1, color=RGBColor(74, 129, 232)):
		p = doc.add_paragraph()
		run = p.add_run(text)
		run.bold = True
		run.font.size = Pt(20 if level == 1 else 16)
		run.font.color.rgb = color
		if level == 1:
			p.space_after = Pt(12)
		else:
			p.space_before = Pt(6)
		return p

	def subheading(text):
		p = doc.add_paragraph()
		run = p.add_run(text)
		run.bold = True
		run.font.size = Pt(14)
		run.font.color.rgb = RGBColor(74, 129, 232)
		p.space_before = Pt(6)
		p.space_after = Pt(4)
		return p

	# parse markdown-like briefing_text produced from page
	for raw_line in briefing_text.splitlines():
		line = raw_line.rstrip()
		ls = line.lstrip()
		if not ls:
			doc.add_paragraph()
			continue
		if ls.startswith("## "):
			heading(ls[3:], level=1)
			continue
		if ls.startswith("### "):
			subheading(ls[4:])
			continue
		if ls.startswith("#### "):
			subheading(ls[5:])
			continue
		if ls.startswith("- ") or ls.startswith("* "):
			content = ls[2:].strip()
			p = doc.add_paragraph(style="List Bullet")
			add_text_with_links(p, content)
			continue
		else:
			p = doc.add_paragraph()
			add_text_with_links(p, line)

	doc.save(bio)
	return bio.getvalue()
st.set_page_config(page_title="AI Brand Reputation Monitoring", layout="wide")
st.title("AI Brand Reputation Monitoring")
st.markdown(
	"This artifact operationalizes the OODA Loop as an AI early-warning system "
	"for detecting reputational threats."
)
st.markdown(
	"OODA Loop, briefly: a fast decision cycle built to outpace opponents in "
	"volatile environments. The four stages **Observe, Orient, Decide, Act** repeat continuously: "
	"you watch signals, interpret them, choose a response, act, then immediately ingest the new "
	"feedback for the next loop. Use it in crises, competitive markets, or incident response when "
	"speed and adaptation are the main advantage."
)

# Simple password gate (password set via APP_PASSWORD in env or Streamlit secrets)
if APP_PASSWORD:
	if not st.session_state["authed"]:
		pwd_input = st.text_input("Enter access password to use the app", type="password", key="app_pwd")
		if pwd_input == APP_PASSWORD:
			st.session_state["authed"] = True
		elif pwd_input:
			st.error("Incorrect password.")
			st.stop()
		else:
			st.stop()
else:
	st.info("Set APP_PASSWORD in secrets/env to protect this app.")

brand = st.text_input("Brand", value=os.getenv("BRAND", "Apple")).strip()

col1, col2 = st.columns([1, 3])

with col1:
	run = st.button("Run monitoring")

if run:
	if not brand:
		st.error("Inserisci un brand.")
		st.stop()

	# passa il brand come env var (senza rifare la pipeline)
	env = os.environ.copy()
	env["BRAND"] = brand
	env["PYTHONUTF8"] = "1"
	env["PYTHONUNBUFFERED"] = "1"
	env["PYTHONIOENCODING"] = "utf-8"

	status_box = st.empty()
	status_box.write("🔄 Starting…")

	with st.expander("Logs", expanded=False):
		log_box = st.empty()
		log_text = ""

	proc = subprocess.Popen(
		[sys.executable, str(ORCH)],
		cwd=str(PROJECT_ROOT),
		env=env,
		stdout=subprocess.PIPE,
		stderr=subprocess.STDOUT,
		text=True,
		bufsize=1,
		encoding="utf-8",
		errors="replace",
	)

	try:
		for line in proc.stdout:
			ln = line.strip()
			low = ln.lower()
			if "init db" in low:
				status_box.write("🔄 Initializing DB…")
			elif "collect rss" in low:
				status_box.write("🔄 Collecting RSS feeds…")
			elif "export raw" in low:
				status_box.write("🔄 Exporting raw data…")
			elif "orient (ai)" in low:
				status_box.write("🔄 Orienting among the content…")
			elif "decide (ai)" in low:
				status_box.write("🔄 Deciding intent and evaluating actions…")
			elif "act (aggregated)" in low or "act (ai)" in low:
				status_box.write("🔄 Crafting action recommendations…")
			elif "pipeline completed" in low:
				status_box.write("✅ Report completed!")

			log_text += line
			try:
				log_box.code(log_text if log_text else "(no stdout)")
			except Exception:
				# UI/websocket closed: stop streaming and terminate process
				if proc.poll() is None:
					proc.terminate()
				break
	finally:
		if proc.poll() is None:
			proc.terminate()
			try:
				proc.wait(timeout=2)
			except Exception:
				pass
	log_box.code(log_text if log_text else "(no stdout)")
	st.caption(f"Exit code: {proc.returncode}")

	# trova l’ultima run e legge last_report_path.txt
	runs_dir = PROJECT_ROOT / "runs"
	if not runs_dir.exists():
		st.error("Cartella runs non trovata.")
		st.stop()

	run_folders = sorted([p for p in runs_dir.iterdir() if p.is_dir()], key=lambda p: p.stat().st_mtime, reverse=True)
	if not run_folders:
		st.error("Nessuna run trovata.")
		st.stop()

	latest_run = run_folders[0]
	last_path_file = latest_run / "last_report_path.txt"
	if not last_path_file.exists():
		st.error("last_report_path.txt non trovato. Assicurati che orchestrator lo scriva.")
		st.stop()

	report_path = Path(last_path_file.read_text(encoding="utf-8").strip())
	if not report_path.exists():
		st.error(f"Report non trovato: {report_path}")
		st.stop()

	# carica foglio REPORT
	try:
		df = pd.read_excel(report_path, sheet_name="REPORT")
	except Exception as e:
		st.error(f"Errore leggendo Excel/REPORT: {e}")
		st.stop()

	# aggiungi colonna brand (dal campo di input)
	df["brand"] = brand

	# prendiamo fino a 30 item (decide è <=30) ordinati per severity e costruiamo un testo compatto per il briefing
	top_items = df.sort_values("severity", ascending=False).head(30)
	items_text = ""
	for _, row in top_items.iterrows():
		items_text += f"""
TITLE: {row['title']}
URL: {row['url']}
PUBLISHED: {row['published_at']}
SEVERITY: {row['severity']}
RISK: {row['reputational_risk']}
ACTION: {row['recommended_action']}
---
"""

	# helper per caricare ultimi export raw/orient (per OBSERVE/ORIENT totali)
	def load_latest_export(run_dir: Path, prefix: str, ext: str = "json") -> pd.DataFrame | None:
		import json
		files = sorted(run_dir.glob(f"{prefix}_*.{ext}"), key=lambda p: p.stat().st_mtime, reverse=True)
		if not files:
			return None
		fp = files[0]
		try:
			if fp.suffix.lower() == ".json":
				with open(fp, "r", encoding="utf-8") as f:
					data = json.load(f)
				return pd.DataFrame(data)
			else:
				return pd.read_csv(fp)
		except Exception:
			return None

	run_dir = latest_run
	raw_df = load_latest_export(run_dir, "raw", "json")
	if raw_df is None:
		raw_df = load_latest_export(run_dir, "raw", "csv")

	orient_df = load_latest_export(run_dir, "orient", "json")
	if orient_df is None:
		orient_df = load_latest_export(run_dir, "orient", "csv")

	# ----------------- UNIFIED REPORT -----------------
	st.subheader("🧠 Executive OODA Briefing (AI Generated)", anchor=None)

	# OBSERVE
	st.subheader("🔍 OBSERVE", anchor=None)
	try:
		with open(PROJECT_ROOT / "config.yaml", "r", encoding="utf-8") as f:
			cfg = yaml.safe_load(f)
		n_feeds = len(cfg.get("observe", {}).get("rss", {}).get("feeds", []))
	except Exception:
		n_feeds = "N/A"

	# defaults for later use
	low_cnt = med_cnt = high_cnt = 0
	top_cats_list = []

	# usa raw export per OBSERVE se disponibile, altrimenti il report ACT
	def _filter_brand_date(df_in: pd.DataFrame) -> pd.DataFrame:
		df_work = df_in.copy()
		# date fallback: published_at -> created_at
		df_work["published_at_dt"] = pd.to_datetime(
			df_work.get("published_at"),
			errors="coerce",
			utc=True,
			infer_datetime_format=True,
		)
		if df_work["published_at_dt"].isna().all() and "created_at" in df_work.columns:
			df_work["published_at_dt"] = pd.to_datetime(
				df_work["created_at"],
				errors="coerce",
				utc=True,
				infer_datetime_format=True,
			)

		# drop rows without a valid datetime before comparisons
		df_work = df_work.dropna(subset=["published_at_dt"])

		# ensure both sides are timezone-aware UTC
		if df_work["published_at_dt"].dt.tz is None:
			df_work["published_at_dt"] = df_work["published_at_dt"].dt.tz_localize("UTC")

		cutoff_local = pd.Timestamp.now(tz="UTC") - pd.Timedelta(days=10)
		date_mask = df_work["published_at_dt"] >= cutoff_local

		if "brand" in df_work.columns:
			brand_mask = df_work["brand"].fillna("").astype(str).str.lower() == brand.lower()
		else:
			# if brand column missing, do not filter by brand
			brand_mask = True

		return df_work[date_mask & brand_mask]

	if raw_df is not None:
		raw_window = _filter_brand_date(raw_df)
		n_items = len(raw_window)
	else:
		df_window = _filter_brand_date(df)
		n_items = len(df_window)

	st.write(f"Found **{n_items}** recent mentions about **{brand}** across **{n_feeds}** monitored sources (last 10 days).")
	st.caption("Goal: detect early signals of reputational risk in real time.")

	# ORIENT
	st.subheader("🧭 ORIENT", anchor=None)
	if orient_df is not None and "reputational_risk" in orient_df.columns:
		orient_window = _filter_brand_date(orient_df)
		low_cnt = (orient_window["reputational_risk"] == "low").sum()
		med_cnt = (orient_window["reputational_risk"] == "medium").sum()
		high_cnt = (orient_window["reputational_risk"] == "high").sum()

		st.write(f"Out of {len(orient_window)} mentions (brand {brand}, last 10 days):")
		st.markdown(
			f"- **{low_cnt}** classified as *low risk*\n"
			f"- **{med_cnt}** classified as *medium risk*\n"
			f"- **{high_cnt}** classified as *high risk*"
		)

		if "narrative_category" in orient_window.columns:
			top_cats = orient_window["narrative_category"].value_counts().head(3)
			top_cats_list = list(top_cats.items())
			st.markdown("Main emerging narratives:")
			for cat, count in top_cats.items():
				st.write(f"- **{cat}** ({count} mentions)")
	else:
		# fallback ai dati del report ACT
		df["published_at_dt"] = pd.to_datetime(df["published_at"], errors="coerce")
		cutoff = pd.Timestamp.now() - pd.Timedelta(days=10)
		df_window = df[
			(df["brand"].str.lower() == brand.lower())
			& (df["published_at_dt"] >= cutoff)
		]
		if "reputational_risk" in df_window.columns:
			low_cnt = (df_window["reputational_risk"] == "low").sum()
			med_cnt = (df_window["reputational_risk"] == "medium").sum()
			high_cnt = (df_window["reputational_risk"] == "high").sum()
			st.write(f"Out of {len(df_window)} mentions (brand {brand}, last 10 days):")
			st.markdown(
				f"- **{low_cnt}** classified as *low risk*\n"
				f"- **{med_cnt}** classified as *medium risk*\n"
				f"- **{high_cnt}** classified as *high risk*"
			)
		else:
			st.write(f"Out of {len(df_window)} mentions (brand {brand}, last 10 days): risk labels not available.")
		if "narrative_category" in df_window.columns:
			top_cats = df_window["narrative_category"].value_counts().head(3)
			top_cats_list = list(top_cats.items())
			st.markdown("Main emerging narratives:")
			for cat, count in top_cats.items():
				st.write(f"- **{cat}** ({count} mentions)")

	# DECIDE
	st.subheader("⚖️ DECIDE", anchor=None)
	# assicurati che df_window esista per le sezioni successive
	if 'df_window' not in locals():
		df["published_at_dt"] = pd.to_datetime(df["published_at"], errors="coerce")
		cutoff = pd.Timestamp.now() - pd.Timedelta(days=10)
		df_window = df[
			(df["brand"].str.lower() == brand.lower())
			& (df["published_at_dt"] >= cutoff)
		]

	if len(df_window) > 0 and "severity" in df_window.columns:
		top_decide = df_window.sort_values("severity", ascending=False).head(3)
		st.markdown("**Top 3 priority issues:**")
		lines = []
		for _, row in top_decide.iterrows():
			lines.append(
				f"- [{row['title']}]({row['url']}) — severity **{row['severity']}**, intent **{row['intent_framing']}**, urgency **{row['urgency']}**"
			)
		st.markdown("\n".join(lines))

		# Disinformation highlight
		if "fact_check_status" in df_window.columns:
			disinfo_mask = (df_window["fact_check_status"] == "disinformation")
			if hasattr(disinfo_mask, "any") and disinfo_mask.any():
				disinfo_df = df_window[disinfo_mask].sort_values("severity", ascending=False)
				st.markdown(f"**{len(disinfo_df)} occurrences were classified as Disinformation:**")
				dis_lines = []
				for _, row in disinfo_df.iterrows():
					title = row.get("title", "Untitled")
					url = row.get("url", "")
					severity = row.get("severity", "N/A")
					intent = row.get("intent_framing", "N/A")
					urg = row.get("urgency", "N/A")
					link_title = f"[{title}]({url})" if url else title
					dis_lines.append(f"- {link_title} — severity **{severity}**, intent **{intent}**, urgency **{urg}**")
				st.markdown("\n".join(dis_lines))
			else:
				st.markdown("**No disinformation sources detected in this window.**")
		else:
			st.markdown("**No disinformation sources detected in this window.**")

		st.markdown(
			"**Intent framing legend:**\n"
			"- **THREAT**: emerging external attack or high-stakes risk to reputation/business.\n"
			"- **DEFENSE**: contain/mitigate an active issue; clarify facts and reduce spread.\n"
			"- **OPPORTUNITY**: positive/neutral momentum that can be leveraged.\n"
			"- **NEUTRAL**: informational, low-signal items to monitor only.\n"
			"- **NOISE**: irrelevant or duplicated content to discard."
		)
	else:
		st.write("No issues to highlight.")

	# ACT + AI briefing
	st.subheader("🚀 ACT", anchor=None)

	# prepare unified dataframe for later reuse (brief + detailed section)
	def _build_df_all():
		def _maybe(name: str):
			try:
				return eval(name)
			except NameError:
				return None

		candidates = [
			_maybe("df_window"),
			_maybe("orient_window"),
			_maybe("raw_window"),
			_maybe("df"),  # full ACT df
			_maybe("orient_df"),
			_maybe("raw_df"),
		]

		df_all_local = None
		for cand in candidates:
			if cand is None:
				continue
			if len(cand) > 0:
				df_all_local = cand
				break
		if df_all_local is None:
			df_all_local = pd.DataFrame()

		if "severity" in df_all_local.columns:
			df_all_local = df_all_local.sort_values("severity", ascending=False)
		return df_all_local

	df_all = _build_df_all()

	with st.spinner("Generating executive summary..."):
		try:
			briefing = generate_ooda_briefing(brand, items_text)
			st.markdown(briefing)

			# build markdown matching on-page content
			page_md_lines = []
			page_md_lines.append(f"## 🧠 Executive OODA Briefing (AI Generated)")
			page_md_lines.append("")
			page_md_lines.append(f"### 🔍 OBSERVE")
			page_md_lines.append(f"Found **{n_items}** recent mentions about **{brand}** across **{n_feeds}** monitored sources (last 10 days).")
			page_md_lines.append("")
			page_md_lines.append("Goal: detect early signals of reputational risk in real time.")
			page_md_lines.append("")
			page_md_lines.append(f"### 🧭 ORIENT")
			page_md_lines.append(f"Out of {n_items} mentions (brand {brand}, last 10 days):")
			page_md_lines.append(f"- **{low_cnt}** classified as *low risk*")
			page_md_lines.append(f"- **{med_cnt}** classified as *medium risk*")
			page_md_lines.append(f"- **{high_cnt}** classified as *high risk*")
			if top_cats_list:
				page_md_lines.append("")
				page_md_lines.append("Main emerging narratives:")
				for cat, count in top_cats_list:
					page_md_lines.append(f"- **{cat}** ({count} mentions)")
			page_md_lines.append("")
			page_md_lines.append(f"### ⚖️ DECIDE")
			if 'top_decide' in locals() and len(top_decide) > 0:
				page_md_lines.append("Top 3 priority issues:")
				for _, row in top_decide.iterrows():
					page_md_lines.append(
						f"- [{row['title']}]({row['url']}) — severity **{row['severity']}**, intent **{row['intent_framing']}**, urgency **{row['urgency']}**"
					)

				# Disinformation list (for download)
				if 'df_window' in locals() and "fact_check_status" in df_window.columns:
					disinfo_df_dl = df_window[df_window["fact_check_status"] == "disinformation"].sort_values("severity", ascending=False)
					if not disinfo_df_dl.empty:
						page_md_lines.append("")
						page_md_lines.append(f"{len(disinfo_df_dl)} occurrences were classified as Disinformation:")
						for _, row in disinfo_df_dl.iterrows():
							title = row.get("title", "Untitled")
							url = row.get("url", "")
							severity = row.get("severity", "N/A")
							intent = row.get("intent_framing", "N/A")
							urg = row.get("urgency", "N/A")
							link_title = f"[{title}]({url})" if url else title
							page_md_lines.append(f"- {link_title} — severity **{severity}**, intent **{intent}**, urgency **{urg}**")
			else:
				page_md_lines.append("- No priority issues available.")
			page_md_lines.append("")
			page_md_lines.append(f"### 🚀 ACT")
			page_md_lines.append(briefing)

			# Append detailed report (ordered by severity desc)
			page_md_lines.append("")
			page_md_lines.append("### 📋 Detailed Report (top by severity)")
			if df_all is not None and not df_all.empty:
				df_sorted = df_all.copy()
				if "severity" in df_sorted.columns:
					df_sorted = df_sorted.sort_values("severity", ascending=False)
				for _, row in df_sorted.iterrows():
					title = row.get("title", "Untitled")
					url = str(row.get("url", "") or "").strip()
					if url and not url.startswith("http"):
						# avoid malformed hrefs in docx
						url = ""
					severity = row.get("severity", "N/A")
					intent = row.get("intent_framing", "N/A")
					source = row.get("source", row.get("brand", ""))
					source_txt = f" • source: {source}" if source else ""
					link_title = f"[{title}]({url})" if url else title
					page_md_lines.append(f"- {link_title} — severity **{severity}**, intent **{intent}**{source_txt}")
			else:
				page_md_lines.append("- No detailed items available.")

			page_md = "\n".join(page_md_lines)

			docx_bytes = briefing_to_docx(brand, page_md)
			st.session_state["brief_docx"] = docx_bytes
			st.session_state["brief_txt"] = page_md
			st.session_state["brief_brand"] = brand
		except Exception as e:
			st.error(f"Briefing generation failed: {e}")


	st.markdown("---")
	st.subheader("📊 Detailed Report", anchor=None)
	# usa tutti i mention disponibili: preferisci dataset con tutte le colonne chiave
	required_cols = ["published_at", "title", "url", "severity", "reputational_risk", "intent_framing"]
	if "df_all" not in locals():
		df_all = _build_df_all()
	# drop duplicates by title+url to avoid repeated rows in detailed report
	if not df_all.empty and {"title", "url"}.issubset(df_all.columns):
		df_all = df_all.drop_duplicates(subset=["title", "url"], keep="first")
	# ensure required columns exist to avoid empty table
	for col in required_cols:
		if col not in df_all.columns:
			df_all[col] = None
	available = required_cols
	display_df = df_all[available].copy()

	st.dataframe(
		display_df,
		width="stretch",
		hide_index=True,
		column_config={
			"url": st.column_config.LinkColumn("url"),
		},
	)

	with col2:
		pass

# keep download available even after widget re-runs
if st.session_state["brief_docx"] or st.session_state["brief_txt"]:
	def _download_link(label: str, data: bytes | str, filename: str, mime: str) -> str:
		if isinstance(data, str):
			data_bytes = data.encode("utf-8")
		else:
			data_bytes = data
		b64 = base64.b64encode(data_bytes).decode()
		return f'<a href="data:{mime};base64,{b64}" download="{filename}" class="st-emotion-cache-link">{label}</a>'

	st.subheader("Download latest briefing")
	brand_for_name = st.session_state.get("brief_brand", "brand")
	if st.session_state["brief_docx"]:
		link = _download_link(
			"Download full briefing (DOCX)",
			st.session_state["brief_docx"],
			f"{brand_for_name}_executive_briefing.docx",
			"application/vnd.openxmlformats-officedocument.wordprocessingml.document",
		)
		st.markdown(link, unsafe_allow_html=True)

	if st.session_state["brief_txt"]:
		link = _download_link(
			"Download full briefing (TXT)",
			st.session_state["brief_txt"],
			f"{brand_for_name}_executive_briefing.txt",
			"text/plain",
		)
		st.markdown(link, unsafe_allow_html=True)
